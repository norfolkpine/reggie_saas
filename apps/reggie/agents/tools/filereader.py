import io
import json
import mimetypes
import os
import tempfile

from agno.tools import Toolkit

# PDF
try:
    import pypdf

    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False

# CSV/Excel
try:
    import pandas as pd

    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

# DOCX
try:
    import docx

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def detect_file_type(file_name: str, file_type_hint: str | None = None) -> str:
    """
    Detect file type using mimetypes, with fallback to extension or provided hint.
    Returns: 'pdf', 'csv', 'docx', 'json', 'md', 'txt'
    """
    # If file_type_hint is a MIME type, map it directly
    mime_map = {
        "application/pdf": "pdf",
        "text/csv": "csv",
        "application/csv": "csv",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "application/msword": "docx",
        "application/json": "json",
        "text/markdown": "md",
        "text/x-markdown": "md",
        "text/plain": "txt",
    }
    if file_type_hint:
        # If it's a known MIME type, map it
        if file_type_hint in mime_map:
            return mime_map[file_type_hint]
        # If it's a known extension, return as is
        ext_hint = file_type_hint.lower().replace(".", "")
        if ext_hint in ["pdf", "csv", "docx", "json", "md", "markdown", "txt"]:
            return ext_hint if ext_hint != "markdown" else "md"
    # Fallback to mimetypes from file_name
    mime, _ = mimetypes.guess_type(file_name)
    if mime:
        if mime in mime_map:
            return mime_map[mime]
        if mime.startswith("text/"):
            return "txt"
    ext = os.path.splitext(file_name)[1].lower()
    if ext == ".pdf":
        return "pdf"
    if ext == ".csv":
        return "csv"
    if ext == ".docx":
        return "docx"
    if ext == ".json":
        return "json"
    if ext in [".md", ".markdown"]:
        return "md"
    if ext == ".txt":
        return "txt"
    return "txt"


class FileReaderTools(Toolkit):
    def __init__(self):
        super().__init__(name="file_reader_tools")
        self.register(self.read_file)

    def read_file(self, content: bytes, file_type: str | None = None, file_name: str | None = None, max_chars: int = 20000) -> str:
        ext = detect_file_type(file_name or "", file_type)
        print(f"[FileReaderTools] Inputs: file_type={file_type}, file_name={file_name}, first 100 bytes={content[:100]}")
        print(f"[FileReaderTools] Detected file type: {ext} for file_name: {str(file_name)[:100]}")
        try:
            if ext in ("pdf", "application/pdf") and PYPDF_AVAILABLE:
                with tempfile.NamedTemporaryFile(suffix=".pdf", delete=True) as tmp:
                    tmp.write(content)
                    tmp.flush()
                    print(f"[FileReaderTools] PDF temp file path: {tmp.name}")
                    text_content = []
                    with open(tmp.name, "rb") as file:
                        pdf_reader = pypdf.PdfReader(file)
                        for page_num, page in enumerate(pdf_reader.pages, 1):
                            try:
                                page_text = page.extract_text()
                                print(f"[FileReaderTools] Page {page_num} text: {repr(page_text)[:100]}")
                                if page_text and page_text.strip():
                                    text_content.append(f"--- Page {page_num} ---\n{page_text}")
                            except Exception as e:
                                print(f"[FileReaderTools] Error extracting text from page {page_num}: {e}")
                                text_content.append(f"--- Page {page_num} (Error: {e}) ---\n")
                    text = "\n\n".join(text_content) if text_content else "No text content found in PDF"
            elif ext == "csv" and PANDAS_AVAILABLE:
                df = pd.read_csv(io.BytesIO(content))
                text = df.to_string(index=False, max_rows=100)
            elif ext == "txt":
                encodings = ["utf-8", "latin-1", "cp1252", "iso-8859-1"]
                for encoding in encodings:
                    try:
                        text = content.decode(encoding)
                        if text.strip():
                            break
                    except UnicodeDecodeError:
                        continue
                else:
                    text = content.decode("utf-8", errors="replace")
            elif ext == "docx" and DOCX_AVAILABLE:
                with tempfile.NamedTemporaryFile(suffix=".docx", delete=True) as tmp:
                    tmp.write(content)
                    tmp.flush()
                    doc = docx.Document(tmp.name)
                    text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            elif ext == "json":
                data = json.loads(content.decode("utf-8"))
                text = json.dumps(data, indent=2)
            elif ext in ("md", "markdown"):
                text = content.decode("utf-8")
            else:
                text = content.decode("utf-8", errors="replace")
            if len(text) > max_chars:
                return text[:max_chars] + "\n... [truncated]"
            return text
        except Exception as e:
            print(f"[FileReaderTools] Error reading file: {e}")
            return f"Error reading file: {e}"

    def read_file_bytes(
        self,
        file_bytes: bytes,
        file_type: str | None = None,
        file_name: str | None = None,
        max_chars: int = 20000,
    ) -> str:
        ext = detect_file_type(file_name or "", file_type)
        print(f"[FileReaderTools] Inputs: file_type={file_type}, file_name={file_name}, first 100 bytes={file_bytes[:100]}")
        print(f"[FileReaderTools] Detected file type: {ext} for file_name: {str(file_name)[:100]}")
        try:
            if ext in ("pdf", "application/pdf") and PYPDF_AVAILABLE:
                with tempfile.NamedTemporaryFile(suffix=".pdf", delete=True) as tmp:
                    tmp.write(file_bytes)
                    tmp.flush()
                    print(f"[FileReaderTools] PDF temp file path: {tmp.name}")
                    text_content = []
                    with open(tmp.name, "rb") as file:
                        pdf_reader = pypdf.PdfReader(file)
                        for page_num, page in enumerate(pdf_reader.pages, 1):
                            try:
                                page_text = page.extract_text()
                                print(f"[FileReaderTools] Page {page_num} text: {repr(page_text)[:100]}")
                                if page_text and page_text.strip():
                                    text_content.append(f"--- Page {page_num} ---\n{page_text}")
                            except Exception as e:
                                print(f"[FileReaderTools] Error extracting text from page {page_num}: {e}")
                                text_content.append(f"--- Page {page_num} (Error: {e}) ---\n")
                    text = "\n\n".join(text_content) if text_content else "No text content found in PDF"
            elif ext == "csv" and PANDAS_AVAILABLE:
                df = pd.read_csv(io.BytesIO(file_bytes))
                text = df.to_string(index=False, max_rows=100)
            elif ext == "txt":
                encodings = ["utf-8", "latin-1", "cp1252", "iso-8859-1"]
                for encoding in encodings:
                    try:
                        text = file_bytes.decode(encoding)
                        if text.strip():
                            break
                    except UnicodeDecodeError:
                        continue
                else:
                    text = file_bytes.decode("utf-8", errors="replace")
            elif ext == "docx" and DOCX_AVAILABLE:
                with tempfile.NamedTemporaryFile(suffix=".docx", delete=True) as tmp:
                    tmp.write(file_bytes)
                    tmp.flush()
                    doc = docx.Document(tmp.name)
                    text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            elif ext == "json":
                text = file_bytes.decode("utf-8", errors="replace")
                data = json.loads(text)
                text = json.dumps(data, indent=2)
            elif ext in ("md", "markdown"):
                text = file_bytes.decode("utf-8", errors="replace")
            else:
                text = file_bytes.decode("utf-8", errors="replace")
            if len(text) > max_chars:
                return text[:max_chars] + "\n... [truncated]"
            return text
        except Exception as e:
            print(f"[FileReaderTools] Error reading file from bytes: {e}")
            return f"Error reading file from bytes: {e}"
